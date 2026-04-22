"""
docqnatool - Streamlit-based document Q&A assistant.

Supports PDF, DOCX, and TXT uploads with optional OCR. Documents are chunked,
embedded with a local TF-IDF vectorizer, stored in a FAISS index, and queried
via ChatGroq (llama-3.3-70b-versatile).
"""

import os
import io
import hashlib
import streamlit as st
from PIL import Image
import pytesseract
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain_core.embeddings import Embeddings
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np

# ----------------- CONFIG ----------------- #
st.set_page_config(
    page_title="docqnatool tool - Smart Document Assistant",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------------- CONSTANTS ----------------- #
LLM_MODEL_NAME = "llama-3.3-70b-versatile"
LLM_TEMPERATURE = 0
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
SIMILARITY_TOP_K = 5

# ----------------- CUSTOM CSS ----------------- #
st.markdown(
    """
<style>
    /* Header */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }

    /* Feature cards */
    .feature-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        border-left: 4px solid #667eea;
        margin: 1rem 0;
        color: #333;
    }
    .feature-card h3, .feature-card h4 { color: #222; }
    .feature-card p, .feature-card li { color: #444; }

    /* Stats */
    .stats-container {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin: 1rem 0;
    }

    /* Upload area */
    .upload-area {
        border: 2px dashed #667eea;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%);
        margin: 1rem 0;
        color: #333;
    }

    /* Chat container */
    .chat-container {
        background: white;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        padding: 1rem;
        margin: 1rem 0;
        color: #222;
    }

    /* Sidebar styling */
    .sidebar .element-container { margin-bottom: 1rem; }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}

    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
    }

    /* File uploader */
    .stFileUploader > div > div {
        background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%);
        border: 2px dashed #667eea;
        border-radius: 10px;
    }

    /* Messages */
    .stSuccess { background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); border-radius: 10px; }
    .stError { background: linear-gradient(135deg, #fa709a 0%, #fee140 100%); border-radius: 10px; }
    .stInfo { background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%); border-radius: 10px; }

    /* Chat bubbles */
    .stChatMessage.user {
        background: #e0f7fa;
        color: #004d40;
        border-radius: 20px 20px 0 20px;
        padding: 0.8rem 1rem;
        margin: 0.5rem 0;
        max-width: 80%;
        align-self: flex-end;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    .stChatMessage.assistant {
        background: #f1f8e9;
        color: #33691e;
        border-radius: 20px 20px 20px 0;
        padding: 0.8rem 1rem;
        margin: 0.5rem 0;
        max-width: 80%;
        align-self: flex-start;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    .stChatMessage { display: flex; flex-direction: column; margin-bottom: 0.5rem; }
    .stChatMessage[data-testid="stChatMessage-system"] {
        background: #eeeeee; color: #555; font-size: 0.9em; border-radius: 10px; text-align: center;
    }
</style>
""",
    unsafe_allow_html=True,
)

# ----------------- UTILS ----------------- #
def get_file_hash(file_content: bytes) -> str:
    """Return the MD5 hex digest of file_content for duplicate detection."""
    return hashlib.md5(file_content).hexdigest()

def normalize_text(text: str) -> str:
    """Normalize whitespace to avoid duplicate detection issues."""
    return " ".join(text.split())

def ocr_image(image_bytes: bytes) -> str:
    """Run Tesseract OCR on image_bytes and return the extracted text, or '' on failure."""
    try:
        image = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(image)
    except Exception:
        return ""

def extract_text_from_pdf(file_bytes: io.BytesIO, use_ocr: bool = True) -> str:
    """Extract text from a PDF byte stream, optionally OCR-ing embedded images."""
    import fitz
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text()
        if page_text:
            text += page_text + "\n"
        if use_ocr:
            for img_meta in page.get_images(full=True):
                try:
                    # img_meta tuple: (xref, smask, width, height, bpc, colorspace, alt, name, filter, enc)
                    xref = img_meta[0]
                    base_image = doc.extract_image(xref)
                    ocr_text = ocr_image(base_image["image"])
                    if ocr_text.strip() and ocr_text not in text:
                        text += "\n" + ocr_text
                except Exception:
                    continue
    return normalize_text(text)

def extract_text_from_docx(file_bytes: io.BytesIO, use_ocr: bool = True) -> str:
    """Extract paragraph text (and optionally image OCR text) from a DOCX byte stream."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_bytes)
    text = "\n".join(para.text for para in doc.paragraphs)
    if use_ocr:
        try:
            for rel in doc.part.rels.values():
                try:
                    if "image" in rel.target_ref:
                        ocr_text = ocr_image(rel.target_part.blob)
                        if ocr_text.strip() and ocr_text not in text:
                            text += "\n" + ocr_text
                except Exception:
                    continue
        except Exception:
            pass
    return normalize_text(text)

def extract_text_from_txt(file_bytes: io.BytesIO) -> str:
    """Decode a plain-text byte stream as UTF-8, falling back to lossy decoding on errors."""
    raw = file_bytes.read()  # read once; the stream is exhausted after this call
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("utf-8", errors="ignore")
    return normalize_text(text)

# ----------------- TFIDF EMBEDDINGS ----------------- #
class TFIDFEmbeddings(Embeddings):
    """
    LangChain-compatible Embeddings implementation backed by scikit-learn TfidfVectorizer.

    The vectorizer is fit lazily on the first call to embed_documents(). Vectors are
    zero-padded or truncated to exactly `max_features` dimensions so FAISS receives a
    uniform-width matrix.
    """

    def __init__(self, max_features: int = 384):
        self.vectorizer = TfidfVectorizer(
            max_features=max_features,
            stop_words="english",
            ngram_range=(1, 2),
            lowercase=True,
            token_pattern=r"\b[a-zA-Z]{2,}\b",
        )
        self.is_fitted = False
        self.dimension = max_features

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if not self.is_fitted:
            self.vectorizer.fit(texts)
            self.is_fitted = True
        vectors = self.vectorizer.transform(texts).toarray()
        return [self._pad(v) for v in vectors]

    def embed_query(self, text: str) -> List[float]:
        if not self.is_fitted:
            return [0.0] * self.dimension
        vec = self.vectorizer.transform([text]).toarray()[0]
        return self._pad(vec)

    def _pad(self, vector: np.ndarray) -> List[float]:
        if len(vector) < self.dimension:
            vector = np.pad(vector, (0, self.dimension - len(vector)), "constant")
        return vector[: self.dimension].tolist()

# ----------------- DOCUMENT MANAGER ----------------- #
class DocumentManager:
    """
    Manages the full document lifecycle: ingestion, deduplication, chunking,
    vector-index construction, and LLM-backed question answering.

    State is stored on st.session_state so it survives Streamlit reruns.
    """

    def __init__(self):
        self.documents: List[Document] = []
        self.processed_files = {}
        self.embeddings = TFIDFEmbeddings()
        self.vectordb = None
        if not os.environ.get("GROQ_API_KEY"):
            st.warning(
                "GROQ_API_KEY is not set. "
                "Document Q&A will be unavailable until a valid key is provided."
            )
            self.llm = None
        else:
            try:
                self.llm = ChatGroq(model_name=LLM_MODEL_NAME, temperature=LLM_TEMPERATURE)
            except Exception as e:
                st.warning(f"Failed to initialize language model: {e}")
                self.llm = None
        self.prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template="""Use the following context to answer the question comprehensively. If you cannot find the answer in the context, say "I cannot find the answer in the provided documents."

Context:
{context}

Question: {question}

Answer:""",
        )

    def add_file(self, filename: str, content: str, file_hash: str, file_size: int):
        if file_hash in self.processed_files:
            return False, f"File '{filename}' already processed (duplicate content)"
        if not content.strip():
            return False, f"File '{filename}' appears to be empty or unreadable"
        doc = Document(page_content=content, metadata={"source": filename, "file_hash": file_hash, "file_size": file_size})
        self.documents.append(doc)
        self.processed_files[file_hash] = {"name": filename, "size": file_size, "word_count": len(content.split())}
        return True, f"✅ Successfully processed '{filename}' ({len(content.split())} words)"

    def _rebuild_vectordb(self):
        if not self.documents:
            self.vectordb = None
            return
        try:
            splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
            chunks = splitter.split_documents(self.documents)
            all_texts = [c.page_content for c in chunks]
            self.embeddings = TFIDFEmbeddings()
            _ = self.embeddings.embed_documents(all_texts)
            self.vectordb = FAISS.from_documents(chunks, self.embeddings)
        except Exception as e:
            st.error(f"Failed to build vector database: {e}")
            self.vectordb = None

    def answer_question(self, question: str) -> str:
        if not self.documents:
            return "❌ No documents uploaded. Please upload some documents first to ask questions."
        if not self.vectordb:
            return "⚠️ Document search index is not ready. Please try uploading documents again."
        if not self.llm:
            return "❌ Language model is not available. Please check your API configuration."
        try:
            docs = self.vectordb.similarity_search(question, k=SIMILARITY_TOP_K)
            if not docs:
                return "🔍 I cannot find any relevant information in the uploaded documents for your question."
            context = "\n\n".join(
                [f"📄 Source: {d.metadata.get('source','Unknown')}\n{d.page_content}" for d in docs]
            )
            chain = self.prompt_template | self.llm
            return chain.invoke({"context": context, "question": question}).content
        except Exception as e:
            return f"❌ Error processing your question: {str(e)}"

    def get_stats(self):
        total_files = len(self.processed_files)
        total_words = sum(info.get("word_count", 0) for info in self.processed_files.values())
        total_size = sum(info.get("size", 0) for info in self.processed_files.values())
        return {"files": total_files, "words": total_words, "size_mb": round(total_size / (1024 * 1024), 2)}

# ----------------- SESSION STATE ----------------- #
if "doc_manager" not in st.session_state:
    st.session_state.doc_manager = DocumentManager()
if "messages" not in st.session_state:
    st.session_state.messages = []

# ----------------- MAIN APP ----------------- #
st.markdown(
    """
<div class="main-header">
    <h1>🧠 docqnatool - Smart Document Assistant</h1>
    <p style="font-size: 1.2em; margin-top: 1rem; opacity: 0.9;">
        Upload your documents and ask intelligent questions powered by AI
    </p>
</div>
""",
    unsafe_allow_html=True,
)

# Sidebar for document upload
with st.sidebar:
    st.markdown("### 📁 Document Upload")
    uploaded_files = st.file_uploader(
        "Choose your documents", type=["pdf", "docx", "txt"], accept_multiple_files=True, help="Supported formats: PDF, DOCX, TXT"
    )
    use_ocr = st.checkbox("🔍 Enable OCR for images", value=True, help="Extract text from images in documents")
    if st.button("🗑️ Clear All Documents"):
        st.session_state.doc_manager = DocumentManager()
        st.session_state.messages = []
        st.success("All documents cleared!")
        st.rerun()
    stats = st.session_state.doc_manager.get_stats()
    if stats["files"] > 0:
        st.markdown("### 📊 Document Statistics")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 Files", stats["files"])
            st.metric("💾 Size (MB)", stats["size_mb"])
        with col2:
            st.metric("📝 Words", f"{stats['words']:,}")
        st.markdown("### 📋 Processed Files")
        # processed_files is keyed by content hash, so two uploads with the same filename
        # but different content will both appear as separate entries. seen_names prevents
        # the same filename being listed twice in the sidebar.
        seen_names = set()
        for file_info in st.session_state.doc_manager.processed_files.values():
            if file_info["name"] not in seen_names:
                st.markdown(f"• **{file_info['name']}** ({file_info['word_count']:,} words)")
                seen_names.add(file_info["name"])

# Process uploaded files
if uploaded_files:
    progress_bar = st.progress(0)
    status_text = st.empty()
    for i, uploaded_file in enumerate(uploaded_files):
        progress_bar.progress((i + 1) / len(uploaded_files))
        status_text.text(f"Processing {uploaded_file.name}...")
        file_data = uploaded_file.getvalue()
        file_bytes = io.BytesIO(file_data)
        file_bytes.seek(0)
        file_hash = get_file_hash(file_data)
        if file_hash in st.session_state.doc_manager.processed_files:
            continue
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        text_content = ""
        try:
            if file_extension == ".pdf":
                text_content = extract_text_from_pdf(file_bytes, use_ocr)
            elif file_extension == ".docx":
                text_content = extract_text_from_docx(file_bytes, use_ocr)
            elif file_extension == ".txt":
                text_content = extract_text_from_txt(file_bytes)
            success, message = st.session_state.doc_manager.add_file(uploaded_file.name, text_content, file_hash, uploaded_file.size)
            if success:
                st.success(message)
            else:
                st.info(message)
        except Exception as e:
            st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
    with st.spinner("🔄 Building search index..."):
        st.session_state.doc_manager._rebuild_vectordb()
    progress_bar.empty()
    status_text.empty()
    st.success("🎉 All documents processed successfully!")

# Main chat interface
if st.session_state.doc_manager.documents:
    st.markdown("### 💬 Chat with your Documents")
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    if prompt := st.chat_input("Ask me anything about your documents..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)
        with st.chat_message("assistant"):
            with st.spinner("🤔 Thinking..."):
                response = st.session_state.doc_manager.answer_question(prompt)
            st.markdown(response)
        st.session_state.messages.append({"role": "assistant", "content": response})
else:
    st.markdown(
        """
    <div class="feature-card">
        <h3>🚀 Get Started</h3>
        <p>Welcome to docqnatool! Here's how to use this intelligent document assistant:</p>
        <ol>
            <li><strong>Upload Documents:</strong> Use the sidebar to upload PDF, DOCX, or TXT files</li>
            <li><strong>Enable OCR:</strong> Check the OCR option to extract text from images in your documents</li>
            <li><strong>Ask Questions:</strong> Once uploaded, ask any questions about your documents</li>
            <li><strong>Get Smart Answers:</strong> The AI will search through your documents and provide detailed answers</li>
        </ol>
    </div>
    """, unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(
            """
        <div class="feature-card">
            <h4>📄 Multiple Formats</h4>
            <p>Support for PDF, DOCX, and TXT files with intelligent text extraction</p>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(
            """
        <div class="feature-card">
            <h4>🔍 OCR Technology</h4>
            <p>Extract text from images and scanned documents automatically</p>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown(
            """
        <div class="feature-card">
            <h4>🧠 AI-Powered</h4>
            <p>Advanced language model provides context-aware answers to your questions</p>
        </div>
        """, unsafe_allow_html=True)

st.markdown("---")
st.markdown(
    """
<div class="footer-content">
    <p>Like the app? 👍</p>
</div>
""", unsafe_allow_html=True)
